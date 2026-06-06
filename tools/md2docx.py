#!/usr/bin/env python3
"""
md2docx.py — convert a Markdown file to a formatted .docx using python-docx.

Reusable across every phase's documentation. Handles the Markdown subset used in
this project's docs: ATX headings (#..####), paragraphs, **bold** + `inline code`
+ [links](url), bullet/numbered lists, GitHub-style tables, fenced ``` code
blocks, horizontal rules, and blockquotes.

Usage:
    python3 tools/md2docx.py docs/Phase-0-Foundation.md [docs/Phase-0-Foundation.docx]

If the output path is omitted, the input's extension is swapped to .docx.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MONO = "Consolas"
CODE_SHADE = "F2F2F2"
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def shade_paragraph(paragraph, fill):
    """Apply a solid background fill to a paragraph (used for code blocks)."""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_inline(paragraph, text):
    """Render inline **bold**, `code`, and [links](url) into a paragraph."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = MONO
            r.font.size = Pt(9.5)
        else:  # link -> show "text (url)"
            mm = LINK.match(tok)
            label, url = mm.group(1), mm.group(2)
            txt = label if (url.startswith("#") or label == url) else f"{label} ({url})"
            paragraph.add_run(txt)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shade_paragraph(p, CODE_SHADE)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = MONO
        r.font.size = Pt(9)
    return p


def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    body = []
    for row in rows[2:]:  # skip the |---|---| separator
        body.append([c.strip() for c in row.strip().strip("|").split("|")])
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            cell = cells[i]
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], r[i] if i < len(r) else "")


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            j = i + 1
            buf = []
            while j < n and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            add_code_block(doc, buf)
            i = j + 1
            continue

        # table (header line that looks like a table, followed by a separator)
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|?[\s:-]+\|", lines[i + 1]):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            h = doc.add_heading(level=level)
            add_inline(h, m.group(2))
            i += 1
            continue

        # horizontal rule
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            doc.add_paragraph().add_run("").add_break()
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped.lstrip("> ").rstrip())
            i += 1
            continue

        # bullet list
        mb = re.match(r"^[-*]\s+(.*)$", stripped)
        if mb:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, mb.group(1))
            i += 1
            continue

        # numbered list
        mn = re.match(r"^\d+\.\s+(.*)$", stripped)
        if mn:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, mn.group(1))
            i += 1
            continue

        # blank line
        if stripped == "":
            i += 1
            continue

        # paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main():
    if len(sys.argv) < 2:
        print("usage: md2docx.py input.md [output.docx]", file=sys.stderr)
        sys.exit(2)
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    convert(md_path, docx_path)


if __name__ == "__main__":
    main()
